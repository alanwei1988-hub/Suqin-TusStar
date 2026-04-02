#!/usr/bin/env python3
import argparse
import base64
import os
from docx import Document


def parse_args():
    parser = argparse.ArgumentParser(description="Create a .docx document from markdown/text content.")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--source", help="Optional source text/markdown file path")
    parser.add_argument("--content-base64", help="Optional base64-encoded utf-8 content")
    parser.add_argument("--title", help="Optional document title")
    return parser.parse_args()


def load_content(args):
    if args.content_base64:
        return base64.b64decode(args.content_base64).decode("utf-8")

    if args.source:
        with open(args.source, "r", encoding="utf-8") as f:
            return f.read()

    raise ValueError("Either --content-base64 or --source must be provided.")


def write_markdown_to_docx(content: str, output_path: str, title: str = ""):
    doc = Document()

    if title and title.strip():
        doc.add_heading(title.strip(), level=0)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        number_prefix = ""
        idx = 0
        while idx < len(stripped) and stripped[idx].isdigit():
            number_prefix += stripped[idx]
            idx += 1
        if number_prefix and idx + 1 < len(stripped) and stripped[idx:idx + 2] == ". ":
            doc.add_paragraph(stripped[idx + 2 :].strip(), style="List Number")
            continue

        doc.add_paragraph(stripped)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)


def main():
    args = parse_args()
    content = load_content(args)
    write_markdown_to_docx(content, args.output, args.title or "")
    print(args.output)


if __name__ == "__main__":
    main()
